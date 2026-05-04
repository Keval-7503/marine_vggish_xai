"""Generate a formatted Word report from Project_Report.md.

This script intentionally keeps the conversion small and local: it supports the
Markdown constructs used by this project report and embeds the project figures.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
REPORT_MD = ROOT / "Project_Report.md"
OUTPUT_DOCX = ROOT / "Marine_VGGish_XAI_Project_Report.docx"
FALLBACK_DOCX = ROOT / "Marine_VGGish_XAI_Project_Report_Updated_8_XAI_Methods.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    paragraph.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_markdown_runs(paragraph, text: str) -> None:
    """Add runs with minimal bold/inline-code support for this report."""
    paragraph.style.font.name = "Times New Roman"
    pattern = re.compile(r"(\*\*.*?\*\*|`.*?`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            run.font.name = "Times New Roman"
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            run.font.name = "Times New Roman"
        else:
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.font.name = "Times New Roman"


def add_table(document: Document, table_lines: list[str]) -> None:
    rows = []
    for line in table_lines:
        parts = [part.strip() for part in line.strip().strip("|").split("|")]
        rows.append(parts)
    if len(rows) < 2:
        return
    header = rows[0]
    body = rows[2:]
    table = document.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, text in enumerate(header):
        set_cell_text(table.rows[0].cells[idx], text, bold=True)
        set_cell_shading(table.rows[0].cells[idx], "D9EAF7")
    for row in body:
        cells = table.add_row().cells
        for idx, text in enumerate(row[: len(header)]):
            set_cell_text(cells[idx], text)
    document.add_paragraph()


def add_image(document: Document, alt: str, rel_path: str) -> None:
    image_path = ROOT / rel_path
    if not image_path.exists():
        paragraph = document.add_paragraph()
        paragraph.add_run(f"[Missing figure: {rel_path}]").italic = True
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(5.9))


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = document.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11)

    for style_name, size in [
        ("Title", 18),
        ("Heading 1", 15),
        ("Heading 2", 13),
        ("Heading 3", 12),
    ]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True


def add_cover_metadata(document: Document, lines: list[str]) -> int:
    title = lines[0].lstrip("#").strip()
    document.add_heading(title, level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER

    idx = 1
    while idx < len(lines) and not lines[idx].strip():
        idx += 1
    meta_lines = []
    while idx < len(lines) and not lines[idx].startswith("## "):
        clean = lines[idx].strip().rstrip("  ")
        if clean:
            meta_lines.append(clean)
        idx += 1

    for meta in meta_lines:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(meta)
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
    document.add_paragraph()
    return idx


def convert() -> Path:
    text = REPORT_MD.read_text(encoding="utf-8")
    lines = text.splitlines()

    document = Document()
    configure_document(document)
    idx = add_cover_metadata(document, lines)

    pending_table: list[str] = []
    list_counter = 0

    def flush_table() -> None:
        nonlocal pending_table
        if pending_table:
            add_table(document, pending_table)
            pending_table = []

    while idx < len(lines):
        raw = lines[idx]
        line = raw.strip()

        if line.startswith("|") and line.endswith("|"):
            pending_table.append(raw)
            idx += 1
            continue
        flush_table()

        if not line:
            list_counter = 0
            idx += 1
            continue

        image_match = re.match(r"!\[(.*?)\]\((.*?)\)", line)
        if image_match:
            add_image(document, image_match.group(1), image_match.group(2))
            idx += 1
            continue

        if line.startswith("## "):
            document.add_heading(line[3:], level=1)
            list_counter = 0
        elif line.startswith("### "):
            document.add_heading(line[4:], level=2)
            list_counter = 0
        elif re.match(r"^\d+\. ", line):
            content = re.sub(r"^\d+\. ", "", line)
            paragraph = document.add_paragraph(style="List Number")
            add_markdown_runs(paragraph, content)
            list_counter += 1
        elif line.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            add_markdown_runs(paragraph, line[2:])
        elif line.startswith("**Figure "):
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(8)
            add_markdown_runs(paragraph, line)
        else:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.line_spacing = 1.08
            paragraph.paragraph_format.space_after = Pt(6)
            add_markdown_runs(paragraph, line)
        idx += 1

    flush_table()

    for section in document.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.text = "Marine VGGish XAI Project Report"
        footer.runs[0].font.name = "Times New Roman"
        footer.runs[0].font.size = Pt(9)

    document.core_properties.title = "Explaining Marine Mammal Sound Classification with VGGish and Time-Frequency XAI"
    document.core_properties.author = "Keval"
    document.core_properties.subject = "XAI project report"
    try:
        document.save(OUTPUT_DOCX)
        return OUTPUT_DOCX
    except PermissionError:
        document.save(FALLBACK_DOCX)
        return FALLBACK_DOCX


if __name__ == "__main__":
    print(convert())
